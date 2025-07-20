import re
import datetime
from docx import Document
import asyncio
from pagespeed_insights2 import main
from docx.enum.text import WD_COLOR_INDEX

def get_name(url):
    name=url.split('https://')[1].split('.')
    name=name[0]+str(datetime.date.today())
    return name

def parse_markdown_with_code(doc: Document, markdown: str):
    # Extract code blocks and replace them with placeholders
    code_blocks = []
    def replacer(match):
        code_blocks.append(match.group(1).strip())
        return f"[[CODE_BLOCK_{len(code_blocks) - 1}]]"

    markdown = re.sub(r"```(?:\w*\n)?(.*?)```", replacer, markdown, flags=re.DOTALL)

    # Now parse line by line
    for line in markdown.splitlines():
        line = line.strip()
        if not line:
            continue

        # Handle placeholders for code blocks
        if line.startswith("[[CODE_BLOCK_"):
            idx = int(re.search(r"\[\[CODE_BLOCK_(\d+)\]\]", line).group(1))
            add_code_block(doc, code_blocks[idx])
            continue

        # Headings and bullets
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("## ") or line.startswith("#### "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("- "):
            add_formatted_paragraph(doc, line[2:], style='List Bullet')
        else:
            add_formatted_paragraph(doc, line)

def add_formatted_paragraph(doc, text, style=None):
    paragraph = doc.add_paragraph(style=style)

    # Highlight image lines
    if text.startswith("<img"):
        run = paragraph.add_run(text)
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        return

    # Bold pattern (**text**)
    pattern = re.compile(r"\*\*(.*?)\*\*")
    pos = 0
    for match in pattern.finditer(text):
        before = text[pos:match.start()]
        bold_text = match.group(1)

        if before:
            paragraph.add_run(before)
        bold_run = paragraph.add_run(bold_text)
        bold_run.bold = True
        pos = match.end()

    if pos < len(text):
        paragraph.add_run(text[pos:])

def add_code_block(doc, block_text):
    for line in block_text.splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE

# Example usage
def generate_docx_from_advice(advice: str, url: str, output_path: str):
    doc = Document()
    parse_markdown_with_code(doc, advice)
    name = get_name(url)
    doc.save(output_path)
